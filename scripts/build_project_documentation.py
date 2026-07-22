from __future__ import annotations

import re
import sys
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
OUTPUT = DOCS / "PROJECT_DOCUMENTATION.docx"
ASSETS = Path("/tmp/fridge-meal-planner-doc-assets")
SKILL = Path(
    "/Users/guangsizeng/.codex/plugins/cache/openai-primary-runtime/"
    "documents/26.715.12143/skills/documents"
)
sys.path.insert(0, str(SKILL / "scripts"))
from table_geometry import apply_table_geometry  # noqa: E402


FILES = [
    "PRODUCT_AND_BUILD_PLAN.md",
    "TECH_STACK_AND_ARCHITECTURE.md",
    "USER_FLOWS.md",
    "BACKEND_AI_AND_API.md",
    "DATA_SCHEMA.md",
    "IMPLEMENTATION_AUDIT.md",
]

NAVY = "17324D"
BLUE = "2E74B5"
LIGHT_BLUE = "E8EEF5"
PALE = "F4F6F9"
INK = "25313C"
MUTED = "66727D"
GRID = "CAD3DC"
WHITE = "FFFFFF"
ORANGE = "E67E22"
GREEN = "3B7D58"
GOLD = "B77B1B"


def rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def set_font(run, size=10.5, bold=False, italic=False, color=INK, name="Aptos"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = rgb(color)


def set_picture_alt(shape, title: str, description: str):
    props = shape._inline.docPr
    props.set("title", title)
    props.set("descr", description)


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    elem = OxmlElement("w:tblHeader")
    elem.set(qn("w:val"), "true")
    tr_pr.append(elem)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:cantSplit"))


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])
    set_font(run, size=9, color=MUTED)


def configure_styles(doc: Document):
    section = doc.sections[0]
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.18

    for name, size, color, before, after in [
        ("Title", 30, NAVY, 0, 8),
        ("Subtitle", 13, MUTED, 0, 12),
        ("Heading 1", 17, BLUE, 16, 8),
        ("Heading 2", 13.5, BLUE, 12, 6),
        ("Heading 3", 11.5, NAVY, 9, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Aptos Display" if name in {"Title", "Heading 1"} else "Aptos"
        style._element.rPr.rFonts.set(qn("w:ascii"), style.font.name)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), style.font.name)
        style.font.size = Pt(size)
        style.font.bold = name != "Subtitle"
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    caption = doc.styles["Caption"]
    caption.font.name = "Aptos"
    caption.font.size = Pt(8.5)
    caption.font.italic = True
    caption.font.color.rgb = rgb(MUTED)
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(8)


def set_header_footer(section):
    header = section.header
    p = header.paragraphs[0]
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(p.add_run("FRIDGE MEAL PLANNER  /  PROJECT HANDBOOK"), size=8.5, bold=True, color=MUTED)
    footer_paragraph = section.footer.paragraphs[0]
    footer_paragraph.clear()
    add_page_number(footer_paragraph)


def add_numbering(doc: Document):
    numbering = doc.part.numbering_part.element

    def make_abstract(abstract_id: int, ordered: bool):
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), "decimal" if ordered else "bullet")
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), "%1." if ordered else "•")
        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), "tab")
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "540")
        tabs.append(tab)
        indent = OxmlElement("w:ind")
        indent.set(qn("w:left"), "540")
        indent.set(qn("w:hanging"), "270")
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "80")
        spacing.set(qn("w:line"), "280")
        spacing.set(qn("w:lineRule"), "auto")
        p_pr.extend([tabs, indent, spacing])
        lvl.extend([start, num_fmt, lvl_text, suff, p_pr])
        abstract.append(lvl)
        numbering.append(abstract)

    def make_num(num_id: int, abstract_id: int):
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abstract_ref = OxmlElement("w:abstractNumId")
        abstract_ref.set(qn("w:val"), str(abstract_id))
        num.append(abstract_ref)
        numbering.append(num)

    make_abstract(90, False)
    make_abstract(91, True)
    make_num(90, 90)
    make_num(91, 91)
    return 90, 91


def new_number_instance(doc: Document, abstract_id=91, start_at=1):
    numbering = doc.part.numbering_part.element
    existing = [int(n.get(qn("w:numId"))) for n in numbering.findall(qn("w:num"))]
    num_id = max(existing, default=99) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    level_override = OxmlElement("w:lvlOverride")
    level_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), str(start_at))
    level_override.append(start_override)
    num.append(level_override)
    numbering.append(num)
    return num_id


def add_rich_text(paragraph, text: str, size=10.5, color=INK):
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*)")
    cursor = 0
    for match in pattern.finditer(text):
        if match.start() > cursor:
            set_font(paragraph.add_run(text[cursor:match.start()]), size=size, color=color)
        token = match.group(0)
        if token.startswith("**"):
            set_font(paragraph.add_run(token[2:-2]), size=size, bold=True, color=color)
        elif token.startswith("`"):
            set_font(paragraph.add_run(token[1:-1]), size=size - 0.5, color=NAVY, name="Aptos Mono")
        else:
            set_font(paragraph.add_run(token[1:-1]), size=size, italic=True, color=color)
        cursor = match.end()
    if cursor < len(text):
        set_font(paragraph.add_run(text[cursor:]), size=size, color=color)


def add_list(doc, text: str, num_id: int):
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])
    p_pr.append(num_pr)
    add_rich_text(p, text)


def add_table(doc, headers, rows):
    col_count = len(headers)
    if col_count == 2:
        widths = [2300, 7060]
    elif col_count == 3:
        widths = [2600, 1100, 5660]
    elif col_count == 4:
        widths = [1700, 1900, 850, 4910]
    else:
        widths = [9360 // col_count] * col_count
        widths[-1] += 9360 - sum(widths)

    table = doc.add_table(rows=1, cols=col_count)
    table.style = "Table Grid"
    table.autofit = False
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = ""
        shade_cell(cell, LIGHT_BLUE)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_font(p.add_run(header), size=8.5, bold=True, color=NAVY)
    repeat_header(table.rows[0])

    for row_data in rows:
        row = table.add_row()
        prevent_row_split(row)
        for idx, value in enumerate(row_data):
            cell = row.cells[idx]
            cell.text = ""
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            add_rich_text(p, value, size=8.3)
    apply_table_geometry(
        table,
        widths,
        table_width_dxa=9360,
        indent_dxa=120,
        cell_margins_dxa={"top": 80, "bottom": 80, "start": 120, "end": 120},
    )
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def load_font(size: int, bold=False):
    candidates = [
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf" if bold else "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


def node(draw, box, title, lines, accent="#2E74B5", fill="#F4F7FA"):
    x1, y1, x2, y2 = box
    draw.rounded_rectangle(box, radius=14, fill=fill, outline="#A9BBCB", width=2)
    draw.rounded_rectangle((x1, y1, x2, y1 + 38), radius=14, fill=accent)
    draw.rectangle((x1, y1 + 20, x2, y1 + 38), fill=accent)
    draw.text((x1 + 12, y1 + 9), title, fill="white", font=load_font(18, True))
    draw.multiline_text((x1 + 12, y1 + 52), "\n".join(lines), fill="#25313C", font=load_font(15), spacing=5)


def arrow(draw, start, end, label=""):
    draw.line([start, end], fill="#7C97AE", width=4)
    x, y = end
    draw.polygon([(x, y), (x - 13, y - 7), (x - 13, y + 7)], fill="#7C97AE")
    if label:
        mx = (start[0] + end[0]) // 2
        my = (start[1] + end[1]) // 2 - 24
        draw.text((mx - 20, my), label, fill="#5D7182", font=load_font(14, True))


def create_diagrams():
    ASSETS.mkdir(parents=True, exist_ok=True)

    architecture = ASSETS / "architecture.png"
    im = Image.new("RGB", (1700, 650), "white")
    d = ImageDraw.Draw(im)
    d.text((55, 35), "Target application architecture", fill="#17324D", font=load_font(30, True))
    boxes = [
        ((55, 145, 300, 340), "Mobile app", ["React Native + Expo", "Typed API client"]),
        ((380, 145, 660, 340), "C# API", ["Auth + validation", "Application services", "Transactions"]),
        ((750, 90, 1025, 270), "PostgreSQL", ["Authoritative data", "Constraints + audit"]),
        ((750, 340, 1025, 520), "Job queue", ["OCR + matching", "Notifications"]),
        ((1120, 145, 1400, 340), "AI layer", ["Structured output", "Allow-listed tools"]),
        ((1460, 145, 1660, 340), "Providers", ["Model", "OCR", "Storage"]),
    ]
    for box, title, lines in boxes:
        node(d, box, title, lines, accent="#E67E22" if title == "Mobile app" else "#2E74B5")
    arrow(d, (300, 242), (380, 242), "HTTPS")
    arrow(d, (660, 210), (750, 180))
    arrow(d, (660, 280), (750, 420))
    arrow(d, (660, 242), (1120, 242))
    arrow(d, (1400, 242), (1460, 242))
    d.rounded_rectangle((370, 565, 1350, 625), radius=16, fill="#FFF4E8", outline="#E0A35C", width=2)
    d.text((420, 584), "Trust boundary: only the API validates and commits authoritative writes.", fill="#7A4B16", font=load_font(20, True))
    im.save(architecture)

    user_flow = ASSETS / "user-flow.png"
    im = Image.new("RGB", (1600, 850), "white")
    d = ImageDraw.Draw(im)
    d.text((55, 35), "Core weekly planning journey", fill="#17324D", font=load_font(30, True))
    flow = [
        ("Verified fridge", ["Manual entry", "Receipt review"]),
        ("Generate proposal", ["Expiry priority", "Diet + schedule"]),
        ("Review and edit", ["Reasons", "Drag / drop"]),
        ("Accept plan", ["Validate again", "Reserve stock"]),
        ("Shopping list", ["Convert units", "Subtract stock"]),
    ]
    x_positions = [55, 360, 665, 970, 1275]
    for index, ((title, lines), x) in enumerate(zip(flow, x_positions)):
        node(d, (x, 225, x + 245, 440), title, lines, accent="#3B7D58" if index == 3 else "#2E74B5")
        if index < len(flow) - 1:
            arrow(d, (x + 245, 332), (x_positions[index + 1], 332))
    d.rounded_rectangle((360, 560, 1240, 690), radius=18, fill="#F4F6F9", outline="#A9BBCB", width=2)
    d.text((410, 585), "Schedule changes loop back to review", fill="#17324D", font=load_font(22, True))
    d.text((410, 630), "Inventory changes trigger revalidation before acceptance.", fill="#506270", font=load_font(19))
    d.line([(980, 560), (980, 470), (787, 470), (787, 440)], fill="#7C97AE", width=4)
    im.save(user_flow)

    schema = ASSETS / "schema.png"
    im = Image.new("RGB", (1700, 1000), "white")
    d = ImageDraw.Draw(im)
    d.text((55, 30), "Implemented relational schema", fill="#17324D", font=load_font(30, True))
    boxes = {
        "Ingredients": (60, 170, 360, 390),
        "FridgeItems": (510, 90, 820, 310),
        "RecipeIngredients": (510, 380, 820, 620),
        "UnitConversions": (510, 700, 820, 900),
        "Recipes": (970, 380, 1250, 600),
        "MealPlans": (1360, 380, 1640, 600),
        "ShoppingLists": (970, 700, 1250, 900),
        "ShoppingListItems": (1360, 700, 1640, 920),
    }
    contents = {
        "Ingredients": ["PK Id", "Name / Category", "DensityGPerMl?"],
        "FridgeItems": ["PK Id", "FK IngredientId", "Quantity / Unit", "BestBefore / Source"],
        "RecipeIngredients": ["PK Id", "FK RecipeId", "FK IngredientId", "Quantity / Unit"],
        "UnitConversions": ["PK Id", "FK IngredientId?", "From / To / Multiplier"],
        "Recipes": ["PK Id", "Name / instructions", "Servings / prep time"],
        "MealPlans": ["PK Id", "FK RecipeId?", "UserId / Date / MealType"],
        "ShoppingLists": ["PK Id", "UserId / WeekStart", "GeneratedAt"],
        "ShoppingListItems": ["PK Id", "FK ShoppingListId", "Name / quantity / unit", "Purchased"],
    }
    for title, box in boxes.items():
        node(d, box, title, contents[title])
    arrow(d, (360, 230), (510, 200), "1:N")
    arrow(d, (360, 300), (510, 500), "1:N")
    arrow(d, (360, 350), (510, 800), "0:N")
    arrow(d, (820, 500), (970, 500), "N:1")
    arrow(d, (1250, 500), (1360, 500), "1:N")
    arrow(d, (1250, 800), (1360, 800), "1:N")
    im.save(schema)

    def simple_flow(filename, title, steps, accepted_index=None):
        path = ASSETS / filename
        canvas = Image.new("RGB", (1600, 900), "white")
        draw = ImageDraw.Draw(canvas)
        draw.text((55, 35), title, fill="#17324D", font=load_font(30, True))
        positions = [
            (55, 160, 445, 330),
            (605, 160, 995, 330),
            (1155, 160, 1545, 330),
            (1155, 535, 1545, 705),
            (605, 535, 995, 705),
            (55, 535, 445, 705),
        ]
        for i, (step_title, detail) in enumerate(steps):
            accent = "#3B7D58" if accepted_index == i else "#2E74B5"
            node(draw, positions[i], step_title, [detail], accent=accent)
            if i < 2:
                arrow(draw, (positions[i][2], 245), (positions[i + 1][0], 245))
            elif i == 2:
                draw.line([(1350, 330), (1350, 535)], fill="#7C97AE", width=4)
                draw.polygon([(1350, 535), (1343, 522), (1357, 522)], fill="#7C97AE")
            elif i < 5:
                arrow(draw, (positions[i][0], 620), (positions[i + 1][2], 620))
        canvas.save(path)
        return path

    onboarding = simple_flow(
        "onboarding-flow.png",
        "First-time setup",
        [
            ("Open app", "Create account or sign in"),
            ("Household", "Create or join shared space"),
            ("Preferences", "Servings, diet, allergens, schedule"),
            ("Add food", "Receipt scan, manual entry, or later"),
            ("Verify", "Confirm safety and imported data"),
            ("Dashboard", "See fridge and expiry priority"),
        ],
        accepted_index=5,
    )
    receipt = simple_flow(
        "receipt-flow.png",
        "Receipt capture and verification",
        [
            ("Capture", "Photograph or upload receipt"),
            ("OCR job", "Extract merchant, date, and lines"),
            ("Match", "Map food lines with confidence"),
            ("Review", "Correct quantity and expiry"),
            ("Confirm", "User approves verified records"),
            ("Import", "Transactional inventory write"),
        ],
        accepted_index=4,
    )
    shopping = simple_flow(
        "shopping-flow.png",
        "Shopping-list generation",
        [
            ("Accepted plan", "Load scheduled recipe servings"),
            ("Requirements", "Expand recipe ingredients"),
            ("Normalize", "Convert compatible units"),
            ("Subtract", "Use available fridge stock"),
            ("Review", "Round, add, edit, or exclude"),
            ("Shop", "Persist checked state"),
        ],
        accepted_index=5,
    )
    ai_lifecycle = simple_flow(
        "ai-lifecycle.png",
        "Validated AI proposal lifecycle",
        [
            ("Request", "Authenticate and validate constraints"),
            ("Context", "Load minimized verified data"),
            ("Generate", "Versioned prompt and JSON schema"),
            ("Validate", "IDs, allergens, slots, quantities"),
            ("Review", "Store proposal; user edits or accepts"),
            ("Commit", "Transaction writes authoritative plan"),
        ],
        accepted_index=4,
    )
    return [architecture, user_flow, schema, onboarding, receipt, shopping, ai_lifecycle]


def diagram_for(code: str, diagrams, index: int):
    if "Expo mobile app" in code or "Target application architecture" in code:
        return diagrams[0], "Target architecture", "Mobile app, C# API, PostgreSQL, jobs, AI orchestration and external providers."
    if "Create account or sign in" in code:
        return diagrams[3], "First-time setup", "Account, household, preferences, food intake, verification and dashboard entry."
    if "Photograph or upload image" in code:
        return diagrams[4], "Receipt verification flow", "Receipt capture, OCR, ingredient matching, user review and transactional inventory import."
    if "Expand recipe requirements" in code:
        return diagrams[5], "Shopping-list flow", "Accepted meals expand into normalized requirements, stock subtraction and persisted shopping completion."
    if code.lstrip().startswith("sequenceDiagram"):
        return diagrams[6], "Validated AI lifecycle", "Authenticated request, minimized context, structured generation, deterministic validation, review and transactional acceptance."
    if "Generate plan" in code or "Backend loads verified context" in code or "Open Planner" in code:
        return diagrams[1], "Weekly planning flow", "Verified inventory leads to a reviewable AI proposal, an accepted plan and a generated shopping list."
    if "Ingredients ||--o{" in code:
        return diagrams[2], "Implemented entity relationship graph", "Eight implemented PostgreSQL tables and their foreign-key relationships."
    return None


def parse_table(lines, start):
    headers = [c.strip() for c in lines[start].strip().strip("|").split("|")]
    rows = []
    idx = start + 2
    while idx < len(lines) and lines[idx].strip().startswith("|"):
        rows.append([c.strip() for c in lines[idx].strip().strip("|").split("|")])
        idx += 1
    return headers, rows, idx


def render_markdown(doc, path: Path, bullet_id: int, number_id: int, diagrams):
    lines = path.read_text(encoding="utf-8").splitlines()
    idx = 0
    paragraph_buffer = []
    first_heading = True

    def flush():
        nonlocal paragraph_buffer
        if paragraph_buffer:
            p = doc.add_paragraph()
            add_rich_text(p, " ".join(s.strip() for s in paragraph_buffer))
            paragraph_buffer = []

    while idx < len(lines):
        raw = lines[idx]
        stripped = raw.strip()
        ordered_match = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if not stripped:
            flush()
            idx += 1
            continue
        if stripped.startswith("```"):
            flush()
            language = stripped[3:].strip()
            code_lines = []
            idx += 1
            while idx < len(lines) and not lines[idx].strip().startswith("```"):
                code_lines.append(lines[idx])
                idx += 1
            code = "\n".join(code_lines)
            diagram = diagram_for(code, diagrams, idx) if language == "mermaid" else None
            if diagram:
                image_path, title, description = diagram
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                image_width = 6.45 if image_path.name in {"architecture.png", "schema.png"} else 5.85
                shape = p.add_run().add_picture(str(image_path), width=Inches(image_width))
                set_picture_alt(shape, title, description)
                cap = doc.add_paragraph(style="Caption")
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap.add_run(title)
            else:
                for code_line in code_lines:
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.25)
                    p.paragraph_format.space_after = Pt(0)
                    set_font(p.add_run(code_line), size=8, color=NAVY, name="Aptos Mono")
            idx += 1
            continue
        if stripped.startswith("# "):
            flush()
            heading_text = re.sub(r"`([^`]+)`", r"\1", stripped[2:])
            if first_heading:
                doc.add_heading(heading_text, level=1)
                first_heading = False
            else:
                doc.add_heading(heading_text, level=1)
            idx += 1
            continue
        if stripped.startswith("## "):
            flush()
            doc.add_heading(re.sub(r"`([^`]+)`", r"\1", stripped[3:]), level=2)
            idx += 1
            continue
        if stripped.startswith("### "):
            flush()
            doc.add_heading(re.sub(r"`([^`]+)`", r"\1", stripped[4:]), level=3)
            idx += 1
            continue
        if stripped.startswith("|") and idx + 1 < len(lines) and re.match(r"^\s*\|[\s:|-]+\|\s*$", lines[idx + 1]):
            flush()
            headers, rows, idx = parse_table(lines, idx)
            add_table(doc, headers, rows)
            continue
        bullet = re.match(r"^-\s+(.*)$", stripped)
        ordered = ordered_match
        if bullet or ordered:
            flush()
            item_text = bullet.group(1) if bullet else ordered.group(2)
            next_idx = idx + 1
            while next_idx < len(lines):
                continuation = lines[next_idx].strip()
                if not continuation:
                    break
                if (
                    continuation.startswith(("#", "|", "```", "- "))
                    or re.match(r"^\d+\.\s+", continuation)
                ):
                    break
                item_text += " " + continuation
                next_idx += 1
            if ordered:
                list_id = new_number_instance(doc, start_at=int(ordered.group(1)))
            else:
                list_id = bullet_id
            add_list(doc, item_text, list_id)
            idx = next_idx
            continue
        if stripped.startswith("**") and stripped.endswith("**") and len(stripped) < 100:
            flush()
            p = doc.add_paragraph()
            add_rich_text(p, stripped)
            idx += 1
            continue
        paragraph_buffer.append(stripped)
        idx += 1
    flush()


def cover(doc: Document):
    doc.add_paragraph().paragraph_format.space_after = Pt(78)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(kicker.add_run("PRODUCT · ENGINEERING · AI"), size=10, bold=True, color=ORANGE)
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Fridge Meal Planner")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Project documentation and implementation handbook")
    doc.add_paragraph().paragraph_format.space_after = Pt(44)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run("Implementation reviewed 22 July 2026"), size=10.5, bold=True, color=NAVY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run("React Native · ASP.NET Core · PostgreSQL · AI orchestration"), size=10, color=MUTED)

    doc.add_paragraph().paragraph_format.space_after = Pt(65)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(
        p.add_run("A practical reference that distinguishes the checked-in prototype from the recommended production design."),
        size=11,
        italic=True,
        color=MUTED,
    )


def build():
    diagrams = create_diagrams()
    doc = Document()
    configure_styles(doc)
    set_header_footer(doc.sections[0])
    bullet_id, number_id = add_numbering(doc)
    cover(doc)
    doc.add_page_break()

    doc.add_heading("Document map", level=1)
    for item in [
        "Product and build plan",
        "Tech stack and architecture",
        "User flows",
        "Backend, AI, and API",
        "Implemented data schema",
        "Current implementation audit",
    ]:
        add_list(doc, item, bullet_id)
    p = doc.add_paragraph()
    add_rich_text(
        p,
        "Status convention: Implemented is present in the repository; Partial exists but is not complete end-to-end; Planned is target design only.",
    )

    for filename in FILES:
        doc.add_page_break()
        render_markdown(doc, DOCS / filename, bullet_id, number_id, diagrams)

    for section in doc.sections:
        set_header_footer(section)

    core = doc.core_properties
    core.title = "Fridge Meal Planner — Project Documentation"
    core.subject = "Product, architecture, user flow, backend, AI, API, schema, and implementation audit"
    core.author = "Fridge Meal Planner project team"
    core.keywords = "meal planning, fridge inventory, C#, PostgreSQL, AI, OCR"
    core.comments = "Generated from implementation-aligned Markdown documentation."
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
